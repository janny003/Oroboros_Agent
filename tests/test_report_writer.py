import ast
import copy
import json
from pathlib import Path

from docx import Document

from tools.report_writer import write_maintenance_report


def minimal_payload(memory_json: str = "inspection_memory.json") -> dict:
    return {
        "generated_at": "2026-05-28 10:00:00",
        "log_root": "tests/fixtures",
        "model_paths": {
            "anomaly_model": "out/model_factory_check/isolation_forest_anomaly_model.pkl",
            "cause_model": "out/model_factory_check/xgboost_fault_cause_classifier.pkl",
        },
        "summary": {
            "total_logs": 1,
            "fail_candidates": 1,
            "high_risk_count": 1,
        },
        "top_causes": [{"label": "통신", "count": 1}],
        "fail_candidates": [
            {
                "file": "sample_test_log.txt",
                "anomaly": -0.12345,
                "cause": "통신",
                "risk": "HIGH",
                "reason": "본문 Failed 1회",
            }
        ],
        "focus": {
            "file": "sample_test_log.txt",
            "summary_text": "우선점검권고: 케이블조립체 / 고장배제 점검 권고: 동일 조건 재시험",
            "preference_note": "(지속 메모리 반영: 시스템제어기조립체 우선)",
        },
        "memory_json": memory_json,
    }


def test_report_writer_writes_docx_and_exact_json_without_mutating_payload(tmp_path):
    payload = minimal_payload(str(tmp_path / "inspection_memory.json"))
    before = copy.deepcopy(payload)
    out_doc = tmp_path / "report.docx"
    out_json = tmp_path / "report.json"

    write_maintenance_report(payload, out_doc, out_json)

    assert payload == before
    assert out_doc.exists()
    assert json.loads(out_json.read_text(encoding="utf-8")) == before
    assert not Path(payload["memory_json"]).exists()

    doc = Document(out_doc)
    paragraph_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )

    assert "정비 통합 보고서" in paragraph_text
    assert "요약" in paragraph_text
    assert "원인 분류 Top 5" in paragraph_text
    assert "FAIL 후보 및 선정 이유" in paragraph_text
    assert "종합의견" in paragraph_text
    assert "우선점검권고" in paragraph_text
    assert "고장배제 점검 권고" in paragraph_text
    assert "원인 라벨" in table_text
    assert "파일" in table_text
    assert "선정 이유" in table_text


def test_report_writer_handles_missing_focus_without_memory_side_effect(tmp_path):
    payload = minimal_payload(str(tmp_path / "inspection_memory.json"))
    payload["focus"] = None
    out_doc = tmp_path / "report.docx"
    out_json = tmp_path / "report.json"

    write_maintenance_report(payload, out_doc, out_json)

    doc = Document(out_doc)
    paragraph_text = "\n".join(p.text for p in doc.paragraphs)
    assert "focus-log가 지정되지 않았거나" in paragraph_text
    assert not Path(payload["memory_json"]).exists()


def test_report_writer_has_no_analysis_or_memory_update_dependencies():
    source = Path("tools/report_writer.py").read_text(encoding="utf-8")
    tree = ast.parse(source)
    forbidden_names = {
        "build_maintenance_analysis_payload",
        "apply_recommendation_policy",
        "load_inspection_memory",
        "save_inspection_memory",
        "update_memory_with_feedback",
        "parse_feedback",
        "pickle",
        "numpy",
    }

    imported = set()
    called = set()
    for node in ast.walk(tree):
        if isinstance(node, ast.ImportFrom):
            for alias in node.names:
                imported.add(alias.name)
        elif isinstance(node, ast.Import):
            for alias in node.names:
                imported.add(alias.name.split(".")[0])
        elif isinstance(node, ast.Call):
            func = node.func
            if isinstance(func, ast.Name):
                called.add(func.id)
            elif isinstance(func, ast.Attribute):
                called.add(func.attr)

    assert not (forbidden_names & imported)
    assert not (forbidden_names & called)
