from __future__ import annotations

import argparse
import csv
import datetime as dt
import json
import pickle
import re
from pathlib import Path

import numpy as np
from docx import Document

try:
    from recommendation_policy import apply_recommendation_policy
except ImportError:  # pragma: no cover - package import path for tests/tools
    from tools.recommendation_policy import apply_recommendation_policy


def read_log_text(path: Path) -> str:
    b = path.read_bytes()
    for enc in ("utf-8", "cp949", "euc-kr", "latin-1"):
        try:
            return b.decode(enc, errors="ignore")
        except Exception:
            continue
    return ""


def extract_metrics(path: Path) -> dict:
    text = read_log_text(path)
    s = (path.name + "\n" + text).lower()

    meas_vals = [float(x) for x in re.findall(r"meas\[\s*([-+]?\d+(?:\.\d+)?)\s*\]", s)]
    min_vals = [float(x) for x in re.findall(r"min\[\s*([-+]?\d+(?:\.\d+)?)\s*\]", s)]
    max_vals = [float(x) for x in re.findall(r"max\[\s*([-+]?\d+(?:\.\d+)?)\s*\]", s)]

    failed_cnt = len(re.findall(r"=>\s*failed", s))
    passed_cnt = len(re.findall(r"=>\s*passed", s))
    retry_cnt = len(re.findall(r"retry|재시도", s))
    crc_cnt = len(re.findall(r"crc|cable|케이블|ethernet", s))

    voltage = float(np.mean(meas_vals)) if meas_vals else (12.0 if "12v" in s else 5.0)
    current = float(np.mean([abs(v) for v in meas_vals[:20]])) / 100.0 if meas_vals else 0.8

    if min_vals and max_vals and meas_vals:
        n = min(len(min_vals), len(max_vals), len(meas_vals))
        margins = []
        for i in range(n):
            low_gap = abs(meas_vals[i] - min_vals[i])
            high_gap = abs(max_vals[i] - meas_vals[i])
            margins.append(min(low_gap, high_gap))
        response_time_ms = float(np.mean(margins) * 1000.0)
    else:
        response_time_ms = 80.0 + 20.0 * retry_cnt

    fail_count = float(failed_cnt if failed_cnt > 0 else s.count("fail"))
    denom = max(1, failed_cnt + passed_cnt)
    crc_error_rate = min(1.0, float(crc_cnt) / float(denom))
    retry_count = float(retry_cnt)

    return {
        "voltage": voltage,
        "current": current,
        "response_time_ms": response_time_ms,
        "fail_count": fail_count,
        "crc_error_rate": crc_error_rate,
        "retry_count": retry_count,
        "failed_lines": failed_cnt,
        "passed_lines": passed_cnt,
    }


def vec_from_metrics(m: dict) -> list[float]:
    return [
        float(m["voltage"]),
        float(m["current"]),
        float(m["response_time_ms"]),
        float(m["fail_count"]),
        float(m["crc_error_rate"]),
        float(m["retry_count"]),
    ]


def pick_existing(paths: list[Path]) -> Path | None:
    for p in paths:
        if p.exists():
            return p
    return None


def classify_cause(model, x: np.ndarray, class_names: list[str] | None = None) -> str:
    try:
        pred = model.predict(x)
        label = pred[0] if hasattr(pred, "__len__") else pred
        if class_names and isinstance(label, (int, np.integer)) and 0 <= int(label) < len(class_names):
            return str(class_names[int(label)])
        return str(label)
    except Exception:
        return "unknown"


def anomaly_score(model, x: np.ndarray) -> float:
    try:
        if hasattr(model, "decision_function"):
            val = model.decision_function(x)
            return float(val[0])
        pred = model.predict(x)
        return float(pred[0])
    except Exception:
        return 0.0


def long_term_risk(anomaly: float, is_fail: bool) -> str:
    base = 0.6 if is_fail else 0.2
    if anomaly < 0:
        base += 0.25
    if base >= 0.75:
        return "HIGH"
    if base >= 0.45:
        return "MEDIUM"
    return "LOW"


def load_fault_exclusion_rows(csv_path: Path) -> list[dict]:
    if not csv_path.exists():
        return []
    rows: list[dict] = []
    with csv_path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for r in reader:
            rows.append(r)
    return rows


def parse_test_ids_from_text(s: str) -> list[str]:
    ids = re.findall(r"\bT(0[1-9]|10)\b", s, flags=re.IGNORECASE)
    out = []
    for n in ids:
        tid = f"T{n}"
        if tid not in out:
            out.append(tid)
    return out


def load_inspection_memory(path: Path) -> dict:
    if not path.exists():
        return {"history": [], "preferences": {}}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {"history": [], "preferences": {}}


def save_inspection_memory(path: Path, memory: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(memory, ensure_ascii=False, indent=2), encoding="utf-8")


def parse_feedback(feedback: str) -> dict:
    out = {"prefer_first_check": "", "resolved": False, "resolved_item": ""}
    if not feedback:
        return out

    m_first = re.search(r"(?:우선점검|먼저점검|first)\s*[:=]\s*([^,;\n]+)", feedback, flags=re.IGNORECASE)
    if m_first:
        out["prefer_first_check"] = m_first.group(1).strip()

    m_resolved = re.search(r"(?:해결|resolved|조치완료)\s*[:=]?\s*(.*)", feedback, flags=re.IGNORECASE)
    if m_resolved:
        out["resolved"] = True
        out["resolved_item"] = (m_resolved.group(1) or "").strip()
    elif any(k in feedback.lower() for k in ["해결", "resolved", "조치완료"]):
        out["resolved"] = True

    return out


def update_memory_with_feedback(memory: dict, feedback: str, focus_name: str, cause: str, risk: str, similar_tests: list[str]) -> None:
    prefs = memory.setdefault("preferences", {})
    hist = memory.setdefault("history", [])

    parsed = parse_feedback(feedback)
    first = str(parsed.get("prefer_first_check", "")).strip()
    if first:
        prefs["prefer_first_check"] = first

    if parsed.get("resolved"):
        solved_map = memory.setdefault("resolved_priority", {})
        key_tests = similar_tests if similar_tests else ["GLOBAL"]
        target = parsed.get("resolved_item") or first or "미기재"
        target = str(target).strip()
        if target:
            for t in key_tests[:5]:
                bucket = solved_map.setdefault(t, {})
                bucket[target] = int(bucket.get(target, 0)) + 1

    hist.append({
        "ts": f"{dt.datetime.now():%Y-%m-%d %H:%M:%S}",
        "focus_log": focus_name,
        "cause": cause,
        "risk": risk,
        "feedback": feedback,
        "prefer_first_check": first,
        "resolved": bool(parsed.get("resolved")),
        "resolved_item": str(parsed.get("resolved_item", "")),
        "similar_tests": similar_tests,
    })
    if len(hist) > 200:
        del hist[:-200]


def cosine_sim(a: list[float], b: list[float]) -> float:
    va = np.array(a, dtype=float)
    vb = np.array(b, dtype=float)
    na = float(np.linalg.norm(va))
    nb = float(np.linalg.norm(vb))
    if na <= 1e-12 or nb <= 1e-12:
        return 0.0
    return float(np.dot(va, vb) / (na * nb))


def guess_test_ids_from_filename(name: str) -> list[str]:
    # 예: "6. RICA_251106.TXT" -> T06
    out = []
    m = re.match(r"\s*(\d{1,2})\.", name)
    if m:
        n = int(m.group(1))
        if 1 <= n <= 10:
            out.append(f"T{n:02d}")
    return out


def collect_similar_cases(focus_row: dict, rows: list[dict], topk: int = 5) -> list[dict]:
    sims = []
    fv = focus_row.get("features") or []
    for r in rows:
        if r is focus_row:
            continue
        score = cosine_sim(fv, r.get("features") or []) if fv else 0.0
        # 동일 시험번호 가중치
        f_t = set(focus_row.get("test_ids", []))
        r_t = set(r.get("test_ids", []))
        if f_t and r_t and (f_t & r_t):
            score += 0.08
        sims.append((score, r))
    sims.sort(key=lambda x: x[0], reverse=True)
    return [{"score": s, **rr} for s, rr in sims[:topk]]


def build_interview_memory_note(memory: dict) -> str:
    last = memory.get("last_interview", {}) if isinstance(memory, dict) else {}
    if not isinstance(last, dict):
        return ""
    answers = last.get("answers", [])
    if not isinstance(answers, list) or not answers:
        return ""

    labels = [
        "1순위 정비대상 확정",
        "동일 조건 재시험",
        "전원/케이블/통신 우선점검 유지",
        "회차별 고위험/기준 적용",
    ]
    pairs = []
    for i, ans in enumerate(answers[:4]):
        label = labels[i] if i < len(labels) else f"질문{i + 1}"
        pairs.append(f"{label}={ans}")
    return "이전 인터뷰 답변 반영: " + ", ".join(pairs)


def apply_interview_priority(memory: dict, exclusion_items: list[str]) -> tuple[list[str], str]:
    last = memory.get("last_interview", {}) if isinstance(memory, dict) else {}
    answers = last.get("answers", []) if isinstance(last, dict) else []
    if not isinstance(answers, list) or len(answers) < 3:
        return exclusion_items, ""

    # Q3: 전원/케이블/통신 라인을 우선 점검 순서로 유지할지 여부.
    # Yes면 다음 진단 권고 앞쪽에 명시적으로 반영한다.
    if str(answers[2]).strip() != "예":
        return exclusion_items, ""

    preferred = "전원/케이블/통신 라인"
    reordered = [preferred]
    for item in exclusion_items:
        if item != preferred and item not in reordered:
            reordered.append(item)
    return reordered[:5], "(이전 인터뷰 답변 반영: 전원/케이블/통신 라인 우선 유지)"


def apply_resolved_priority(memory: dict, test_ids: list[str], exclusion_items: list[str]) -> tuple[list[str], str]:
    solved_map = memory.get("resolved_priority", {}) if isinstance(memory, dict) else {}
    merged = {}
    for t in (test_ids or []):
        for k, v in (solved_map.get(t, {}) or {}).items():
            merged[k] = merged.get(k, 0) + int(v)
    for k, v in (solved_map.get("GLOBAL", {}) or {}).items():
        merged[k] = merged.get(k, 0) + int(v)

    if not merged:
        return exclusion_items, ""

    ranked = [k for k, _ in sorted(merged.items(), key=lambda kv: kv[1], reverse=True)]
    reordered = []
    used = set()

    for item in ranked:
        if item not in used:
            reordered.append(item)
            used.add(item)
    for item in exclusion_items:
        if item not in used:
            reordered.append(item)
            used.add(item)

    note = f"(지속 메모리 반영: 과거 해결 이력 기반 우선항목 {min(3, len(ranked))}개를 앞에 배치)"
    return reordered[:5], note


def build_exclusion_recommendation(log_file: Path, log_text: str, fault_rows: list[dict], cause: str) -> tuple[list[str], str, list[str]]:
    if not fault_rows:
        items = ["통신 경로", "전원 경로", "케이블/커넥터"]
        checklist = (
            "고장배제 점검 권고: ① 통신(이더넷 링크/포트) 확인, "
            "② 전원(220VAC/28V) 안정성 확인, ③ 케이블/커넥터 접촉 재체결 후 재시험"
        )
        return items, checklist, ["GLOBAL"]

    test_ids = parse_test_ids_from_text(log_file.name + "\n" + log_text)
    matched = [r for r in fault_rows if (r.get("test_id") or "") in test_ids]

    # 테스트 ID 매칭이 없으면 원인 라벨/로그 키워드 기반으로 전역 우선순위 선정
    if not matched:
        key = (cause or "").lower() + " " + log_text.lower()
        if any(k in key for k in ["rf", "freq", "channel"]):
            matched = [r for r in fault_rows if (r.get("test_id") or "") in {"T08", "T09"}]
        elif any(k in key for k in ["ethernet", "port", "crc", "retry", "통신"]):
            matched = [r for r in fault_rows if (r.get("test_id") or "") in {"T03", "T04", "T05", "T06", "T10"}]
        elif any(k in key for k in ["power", "28v", "220vac", "전원"]):
            matched = [r for r in fault_rows if (r.get("test_id") or "") in {"T01", "T02"}]

    guides: list[str] = []
    for r in matched:
        g = (r.get("field_guidance") or "").strip()
        if g and g not in guides and "후속 분석 필요" not in g:
            guides.append(g)
        if len(guides) >= 3:
            break

    if not guides:
        guides = ["시스템제어기조립체 우선 점검", "케이블조립체(TW605/TW606) 우선 점검", "전원제어기 경로 우선 점검"]

    checklist = (
        "고장배제 점검 권고: ① 통신 경로(이더넷/포트/CRC·retry) 확인, "
        "② 전원 경로(입력전원/28V 출력) 확인, ③ 제어기 간 연동 상태 확인, "
        "④ 동일 조건 재시험으로 재현성 검증"
    )
    if not test_ids:
        test_ids = guess_test_ids_from_filename(log_file.name)
    if not test_ids:
        test_ids = ["GLOBAL"]

    return guides, checklist, test_ids


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--log-root", required=True)
    ap.add_argument("--out-doc", required=True)
    ap.add_argument("--out-json", default="", help="추가 저장용 JSON 보고서 경로 (미지정 시 out-doc와 같은 이름의 .json)")
    ap.add_argument("--project-root", required=True)
    ap.add_argument("--focus-log", default="", help="특정 시험 로그 파일 경로 (예: .../6. RICA_251106.TXT)")
    ap.add_argument("--operator-feedback", default="", help="운용자 입력 메모. 예) 우선점검: 시스템제어기조립체")
    ap.add_argument("--memory-json", default="", help="지속 점검 메모리 JSON 경로")
    ap.add_argument(
        "--fault-exclusion-csv",
        default="C:/Users/yjs/Desktop/JAN/Policy/OrobrosTest/data/fault_exclusion_master_map.csv",
        help="시험 항목별 고장배제 매트릭스 CSV 경로",
    )
    args = ap.parse_args()

    log_root = Path(args.log_root)
    out_doc = Path(args.out_doc)
    out_json = Path(args.out_json) if args.out_json else out_doc.with_suffix(".json")
    project_root = Path(args.project_root)

    if not log_root.exists():
        raise SystemExit(f"log root not found: {log_root}")

    iso_path = pick_existing([
        project_root / "out" / "model_factory_check_xgb" / "isolation_forest_anomaly_model.pkl",
        project_root / "out" / "model_factory_check" / "isolation_forest_anomaly_model.pkl",
    ])
    cause_path = pick_existing([
        project_root / "out" / "model_factory_check_xgb" / "xgboost_fault_cause_classifier.pkl",
        project_root / "out" / "model_factory_check" / "xgboost_fault_cause_classifier.pkl",
    ])

    if not iso_path or not cause_path:
        raise SystemExit("required models not found in out/model_factory_check*")

    iso_obj = pickle.load(open(iso_path, "rb"))
    cause_obj = pickle.load(open(cause_path, "rb"))
    iso_model = iso_obj.get("model", iso_obj) if isinstance(iso_obj, dict) else iso_obj
    cause_model = cause_obj.get("model", cause_obj) if isinstance(cause_obj, dict) else cause_obj
    class_names = cause_obj.get("classes") if isinstance(cause_obj, dict) else None

    logs = sorted(set(list(log_root.rglob("*.TXT")) + list(log_root.rglob("*.txt"))))
    if not logs:
        raise SystemExit("no TXT logs found")

    rows = []
    for p in logs:
        m = extract_metrics(p)
        f = vec_from_metrics(m)
        x = np.array([f], dtype=float)
        a = anomaly_score(iso_model, x)
        is_fail = (m["failed_lines"] > 0) or ("fail" in p.name.lower())
        cause = classify_cause(cause_model, x, class_names)
        risk = long_term_risk(a, is_fail)
        reason_bits = []
        if m["failed_lines"] > 0:
            reason_bits.append(f"본문 Failed {m['failed_lines']}회")
        if "fail" in p.name.lower():
            reason_bits.append("파일명에 FAIL 포함")
        if a < 0:
            reason_bits.append("이상탐지 점수<0")
        if f[5] > 0:
            reason_bits.append(f"retry 키워드 {int(f[5])}회")
        reason = ", ".join(reason_bits) if reason_bits else "정상 패턴 우세"
        test_ids = parse_test_ids_from_text(p.name + "\n" + read_log_text(p))
        if not test_ids:
            test_ids = guess_test_ids_from_filename(p.name)
        rows.append({
            "file": str(p),
            "anomaly": a,
            "cause": cause,
            "risk": risk,
            "is_fail": is_fail,
            "reason": reason,
            "features": f,
            "test_ids": test_ids,
        })

    total = len(rows)
    fail_n = sum(1 for r in rows if r["is_fail"])
    high_n = sum(1 for r in rows if r["risk"] == "HIGH")

    cause_counts = {}
    for r in rows:
        cause_counts[r["cause"]] = cause_counts.get(r["cause"], 0) + 1

    filtered_causes = [(k, v) for k, v in cause_counts.items() if str(k).strip().lower() != "normal"]
    top_causes = sorted(filtered_causes, key=lambda kv: kv[1], reverse=True)[:5]
    if not top_causes:
        top_causes = [("(normal 제외 후 항목 없음)", 0)]

    focus_log = Path(args.focus_log) if args.focus_log else None
    focus_row = None
    if focus_log:
        f_resolved = focus_log.resolve()
        for r in rows:
            rp = Path(r["file"])
            if rp.resolve() == f_resolved:
                focus_row = r
                break
        if focus_row is None:
            # 경로가 일부 생략된 경우(예: 하위 폴더 누락), 파일명 기준으로 재매칭
            for r in rows:
                if Path(r["file"]).name.lower() == focus_log.name.lower():
                    focus_row = r
                    break

    fault_rows = load_fault_exclusion_rows(Path(args.fault_exclusion_csv))
    memory_path = Path(args.memory_json) if args.memory_json else (project_root / "out" / "inspection_memory.json")
    memory = load_inspection_memory(memory_path)

    fail_rows = [r for r in rows if r["is_fail"]]

    doc = Document()
    report_payload: dict = {
        "generated_at": f"{dt.datetime.now():%Y-%m-%d %H:%M:%S}",
        "log_root": str(log_root),
        "model_paths": {
            "anomaly_model": str(iso_path),
            "cause_model": str(cause_path),
        },
        "summary": {
            "total_logs": total,
            "fail_candidates": fail_n,
            "high_risk_count": high_n,
        },
        "top_causes": [{"label": str(k), "count": int(v)} for k, v in top_causes],
        "fail_candidates": [
            {
                "file": Path(r["file"]).name,
                "anomaly": float(r["anomaly"]),
                "cause": str(r["cause"]),
                "risk": str(r["risk"]),
                "reason": str(r["reason"]),
            }
            for r in fail_rows[:80]
        ],
        "focus": None,
        "memory_json": str(memory_path),
    }
    doc.add_heading("정비 통합 보고서", level=1)
    doc.add_paragraph(f"생성시각: {dt.datetime.now():%Y-%m-%d %H:%M:%S}")
    doc.add_paragraph(f"로그 경로: {log_root}")
    doc.add_paragraph(f"사용 이상탐지 모델: {iso_path}")
    doc.add_paragraph(f"사용 원인분류 모델: {cause_path}")

    doc.add_heading("요약", level=2)
    doc.add_paragraph(f"총 로그: {total}건")
    doc.add_paragraph(f"FAIL 후보: {fail_n}건")
    doc.add_paragraph(f"장기 고장 HIGH 위험: {high_n}건")

    doc.add_heading("원인 분류 Top 5", level=2)
    t1 = doc.add_table(rows=1, cols=2)
    t1.rows[0].cells[0].text = "원인 라벨"
    t1.rows[0].cells[1].text = "건수"
    for k, v in top_causes:
        c = t1.add_row().cells
        c[0].text = str(k)
        c[1].text = str(v)

    doc.add_heading("FAIL 후보 및 선정 이유", level=2)
    doc.add_paragraph(f"총 FAIL 후보: {len(fail_rows)}건")
    doc.add_paragraph('이상점수는 Isolation Forest가 주는 "정상에서 얼마나 벗어났는지" 점수입니다.')
    t3 = doc.add_table(rows=1, cols=5)
    t3.rows[0].cells[0].text = "파일"
    t3.rows[0].cells[1].text = "이상점수"
    t3.rows[0].cells[2].text = "원인 라벨"
    t3.rows[0].cells[3].text = "위험도"
    t3.rows[0].cells[4].text = "선정 이유"
    for r in fail_rows[:80]:
        c = t3.add_row().cells
        c[0].text = Path(r["file"]).name
        c[1].text = f"{r['anomaly']:.4f}"
        c[2].text = str(r["cause"])
        c[3].text = r["risk"]
        c[4].text = r["reason"]

    doc.add_heading("종합의견", level=2)
    if focus_row:
        focus_path = Path(focus_row["file"])
        focus_text = read_log_text(focus_path)
        hist = [r for r in rows if Path(r["file"]).resolve() != focus_path.resolve()]
        hist_anom = [float(r["anomaly"]) for r in hist] if hist else [0.0]
        hist_fail_rate = (sum(1 for r in hist if r["is_fail"]) / len(hist)) if hist else 0.0

        # short-term anomaly diagnosis
        pctl = 50.0
        if hist_anom:
            pctl = float((sum(1 for v in hist_anom if v <= float(focus_row["anomaly"])) / len(hist_anom)) * 100.0)
        short_diag = (
            f"단기 진단 기준으로 {focus_path.name}의 이상점수는 {focus_row['anomaly']:.4f}이며, "
            f"이전 로그 대비 약 하위 {100.0 - pctl:.1f}% 수준의 편차를 보였습니다. "
            f"원인분류는 '{focus_row['cause']}'로 판정되었습니다."
        )

        # mid/long term forecast sentence
        risk = focus_row["risk"]
        if risk == "HIGH":
            trend = "중장기 고장위험이 높아, 예방정비 주기를 즉시 단축하는 것이 필요합니다."
        elif risk == "MEDIUM":
            trend = "중장기 고장위험이 중간 수준으로, 추세 모니터링과 점검주기 보정이 권고됩니다."
        else:
            trend = "중장기 고장위험은 낮은 편이나, 동일 증상 반복 여부를 주기적으로 추적해야 합니다."

        exclusion_items, exclusion_check, focus_test_ids = build_exclusion_recommendation(
            focus_path, focus_text, fault_rows, str(focus_row["cause"])
        )

        similar_cases = collect_similar_cases(focus_row, rows, topk=5)
        similar_case_line = ""
        if similar_cases:
            names = [Path(r["file"]).name for r in similar_cases[:3]]
            similar_case_line = "유사 이력: " + " / ".join(names)

        recommendation = apply_recommendation_policy(memory, focus_test_ids, exclusion_items)
        exclusion_items = list(recommendation.get("recommended_exclusion_items", exclusion_items))[:3]
        memory_note = str(recommendation.get("resolved_priority_note", ""))
        interview_priority_note = str(recommendation.get("interview_priority_note", ""))
        interview_memory_note = str(recommendation.get("interview_memory_note", ""))
        preference_note = str(recommendation.get("preference_note", ""))
        exclusion_line = "우선점검권고: " + " / ".join(exclusion_items)

        summary = (
            f"본 시험 로그({focus_path.name})를 이전 누적 데이터({len(hist)}건)와 비교 분석한 결과, "
            f"이전 FAIL 비율은 {hist_fail_rate*100.0:.1f}%였습니다. {short_diag} "
            f"중장기 예측 위험도는 {risk}로 평가되었고, {trend} {similar_case_line} "
            f"{exclusion_line} {exclusion_check} {memory_note} {interview_priority_note} {interview_memory_note}"
        )
        doc.add_paragraph(summary)
        report_payload["focus"] = {
            "file": focus_path.name,
            "anomaly": float(focus_row["anomaly"]),
            "cause": str(focus_row["cause"]),
            "risk": str(risk),
            "short_diag": short_diag,
            "similar_case_names": [Path(r["file"]).name for r in similar_cases[:3]],
            "recommended_exclusion_items": exclusion_items,
            "checklist": exclusion_check,
            "summary_text": summary,
            "test_ids": focus_test_ids,
            "interview_memory_note": interview_memory_note,
            "interview_priority_note": interview_priority_note,
        }

        update_memory_with_feedback(
            memory,
            args.operator_feedback,
            focus_path.name,
            str(focus_row["cause"]),
            str(risk),
            focus_test_ids,
        )
        save_inspection_memory(memory_path, memory)
        if args.operator_feedback:
            doc.add_paragraph(f"운용자 피드백 반영: {args.operator_feedback}")
        if preference_note:
            doc.add_paragraph(preference_note)
    else:
        doc.add_paragraph(
            "focus-log가 지정되지 않았거나 로그 루트 내에서 찾지 못해, 본 보고서는 전체 집계 기반 종합의견만 제공합니다."
        )

    out_doc.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_doc)
    out_json.parent.mkdir(parents=True, exist_ok=True)
    out_json.write_text(json.dumps(report_payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(str(out_doc))
    print(str(out_json))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
