from __future__ import annotations

import json
from pathlib import Path

from docx import Document


def write_maintenance_report(
    payload: dict,
    out_doc: Path | str,
    out_json: Path | str,
    operator_feedback: str = "",
) -> None:
    """Serialize a maintenance report payload to DOCX and JSON.

    This writer intentionally performs no analysis, recommendation, or memory
    updates. It renders the supplied payload as-is and writes the same payload to
    JSON with Korean-safe UTF-8 settings.
    """
    out_doc = Path(out_doc)
    out_json = Path(out_json)

    summary = payload["summary"]
    top_causes = [(r["label"], r["count"]) for r in payload["top_causes"]]
    fail_candidates = payload["fail_candidates"]
    focus = payload.get("focus")
    iso_path = payload["model_paths"]["anomaly_model"]
    cause_path = payload["model_paths"]["cause_model"]

    doc = Document()
    doc.add_heading("정비 통합 보고서", level=1)
    doc.add_paragraph(f"생성시각: {payload['generated_at']}")
    doc.add_paragraph(f"로그 경로: {payload['log_root']}")
    doc.add_paragraph(f"사용 이상탐지 모델: {iso_path}")
    doc.add_paragraph(f"사용 원인분류 모델: {cause_path}")

    doc.add_heading("요약", level=2)
    doc.add_paragraph(f"총 로그: {summary['total_logs']}건")
    doc.add_paragraph(f"FAIL 후보: {summary['fail_candidates']}건")
    doc.add_paragraph(f"장기 고장 HIGH 위험: {summary['high_risk_count']}건")

    doc.add_heading("원인 분류 Top 5", level=2)
    t1 = doc.add_table(rows=1, cols=2)
    t1.rows[0].cells[0].text = "원인 라벨"
    t1.rows[0].cells[1].text = "건수"
    for label, count in top_causes:
        cells = t1.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(count)

    doc.add_heading("FAIL 후보 및 선정 이유", level=2)
    doc.add_paragraph(f"총 FAIL 후보: {summary['fail_candidates']}건")
    doc.add_paragraph('이상점수는 Isolation Forest가 주는 "정상에서 얼마나 벗어났는지" 점수입니다.')
    t3 = doc.add_table(rows=1, cols=5)
    t3.rows[0].cells[0].text = "파일"
    t3.rows[0].cells[1].text = "이상점수"
    t3.rows[0].cells[2].text = "원인 라벨"
    t3.rows[0].cells[3].text = "위험도"
    t3.rows[0].cells[4].text = "선정 이유"
    for row in fail_candidates:
        cells = t3.add_row().cells
        cells[0].text = str(row["file"])
        cells[1].text = f"{float(row['anomaly']):.4f}"
        cells[2].text = str(row["cause"])
        cells[3].text = str(row["risk"])
        cells[4].text = str(row["reason"])

    doc.add_heading("종합의견", level=2)
    if focus:
        doc.add_paragraph(str(focus["summary_text"]))
        if operator_feedback:
            doc.add_paragraph(f"운용자 피드백 반영: {operator_feedback}")
        preference_note = str(focus.get("preference_note", ""))
        if preference_note:
            doc.add_paragraph(preference_note)
    else:
        doc.add_paragraph(
            "focus-log가 지정되지 않았거나 로그 루트 내에서 찾지 못해, 본 보고서는 전체 집계 기반 종합의견만 제공합니다."
        )

    out_doc.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_doc)
    out_json.parent.mkdir(parents=True, exist_ok=True)
    out_json.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
